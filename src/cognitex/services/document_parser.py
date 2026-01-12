"""Local document parsing using python-docx, PyPDF2, pdfplumber, and openpyxl.

This module provides fallback document analysis when Anthropic Skills are unavailable.
It extracts text, tracked changes, comments, and highlights from documents.
"""

import io
from dataclasses import dataclass, field
from typing import Any

import structlog

logger = structlog.get_logger()


@dataclass
class DocumentAnalysis:
    """Result of document analysis."""

    summary: str = ""
    changes: list[str] = field(default_factory=list)
    review_items: list[str] = field(default_factory=list)
    questions: list[dict[str, str]] = field(default_factory=list)
    filename: str = ""
    method: str = "local"  # "local" or "skills"
    raw_text: str = ""

    def to_dict(self) -> dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return {
            "summary": self.summary,
            "changes": self.changes,
            "review_items": self.review_items,
            "questions": self.questions,
            "filename": self.filename,
            "method": self.method,
        }


class UnsupportedDocumentType(Exception):
    """Raised when document type is not supported for parsing."""

    pass


class LocalDocumentParser:
    """Fallback document parsing using local Python libraries.

    Supports:
    - DOCX: python-docx for text, tracked changes, comments, highlights
    - PDF: pdfplumber for text extraction
    - XLSX: openpyxl for spreadsheet analysis
    """

    # MIME type to extension mapping
    MIME_TO_EXT = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "application/vnd.ms-excel": "xls",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation": "pptx",
    }

    SUPPORTED_TYPES = {"docx", "pdf", "xlsx"}

    def get_extension(self, filename: str, mime_type: str | None = None) -> str:
        """Get file extension from filename or MIME type."""
        if mime_type and mime_type in self.MIME_TO_EXT:
            return self.MIME_TO_EXT[mime_type]

        if "." in filename:
            return filename.rsplit(".", 1)[-1].lower()

        return ""

    def is_supported(self, filename: str, mime_type: str | None = None) -> bool:
        """Check if document type is supported."""
        ext = self.get_extension(filename, mime_type)
        return ext in self.SUPPORTED_TYPES

    async def parse(
        self,
        filename: str,
        content: bytes,
        mime_type: str | None = None,
    ) -> DocumentAnalysis:
        """
        Parse a document and extract analysis.

        Args:
            filename: Name of the file
            content: Raw file bytes
            mime_type: Optional MIME type for better detection

        Returns:
            DocumentAnalysis with extracted content

        Raises:
            UnsupportedDocumentType: If file type not supported
        """
        ext = self.get_extension(filename, mime_type)

        if ext == "docx":
            return await self._parse_docx(filename, content)
        elif ext == "pdf":
            return await self._parse_pdf(filename, content)
        elif ext == "xlsx":
            return await self._parse_xlsx(filename, content)
        else:
            raise UnsupportedDocumentType(
                f"Cannot parse '{filename}' with extension '{ext}'. "
                f"Supported types: {', '.join(self.SUPPORTED_TYPES)}"
            )

    async def _parse_docx(self, filename: str, content: bytes) -> DocumentAnalysis:
        """Parse DOCX file using python-docx.

        Extracts:
        - Full text content
        - Tracked changes (insertions/deletions)
        - Comments
        - Highlighted text
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            logger.error("python-docx not installed, cannot parse DOCX")
            return DocumentAnalysis(
                filename=filename,
                summary="[Error: python-docx library not installed]",
                method="local",
            )

        try:
            doc = Document(io.BytesIO(content))

            # Extract all text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            full_text = "\n".join(paragraphs)

            # Find tracked changes (revisions)
            changes = []
            for para in doc.paragraphs:
                for run in para.runs:
                    element = run._element

                    # Check for insertions (w:ins elements)
                    ins_elements = element.xpath(".//w:ins", namespaces=element.nsmap)
                    if ins_elements:
                        for ins in ins_elements:
                            # Get the inserted text
                            text = "".join(t.text or "" for t in ins.iter() if hasattr(t, "text"))
                            if text.strip():
                                changes.append(f"Added: {text.strip()}")

                    # Check for deletions (w:del elements)
                    del_elements = element.xpath(".//w:del", namespaces=element.nsmap)
                    if del_elements:
                        for del_elem in del_elements:
                            text = "".join(t.text or "" for t in del_elem.iter() if hasattr(t, "text"))
                            if text.strip():
                                changes.append(f"Deleted: {text.strip()}")

            # Find comments
            questions = []
            try:
                # Access comments through the document part
                comments_part = doc.part.comments_part
                if comments_part:
                    comments_element = comments_part.element
                    for comment in comments_element.findall(qn("w:comment")):
                        author = comment.get(qn("w:author"), "Unknown")
                        # Extract comment text
                        comment_text = "".join(
                            t.text or ""
                            for t in comment.iter()
                            if hasattr(t, "text") and t.text
                        )
                        if comment_text.strip():
                            questions.append({
                                "author": author,
                                "text": comment_text.strip(),
                            })
            except Exception as e:
                logger.debug("Could not extract comments", error=str(e))

            # Find highlighted text
            review_items = []
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.highlight_color:
                        text = run.text.strip()
                        if text and text not in review_items:
                            review_items.append(text)

            # Create summary from first few paragraphs
            summary_paras = paragraphs[:3] if paragraphs else []
            summary = " ".join(summary_paras)[:500]
            if len(summary) == 500:
                summary = summary[:497] + "..."

            return DocumentAnalysis(
                filename=filename,
                summary=summary,
                changes=changes,
                review_items=review_items,
                questions=questions,
                raw_text=full_text,
                method="local",
            )

        except Exception as e:
            logger.error("Failed to parse DOCX", filename=filename, error=str(e))
            return DocumentAnalysis(
                filename=filename,
                summary=f"[Error parsing DOCX: {str(e)}]",
                method="local",
            )

    async def _parse_pdf(self, filename: str, content: bytes) -> DocumentAnalysis:
        """Parse PDF file using pdfplumber.

        Extracts text content. PDFs don't have tracked changes like DOCX.
        """
        try:
            import pdfplumber
        except ImportError:
            logger.error("pdfplumber not installed, cannot parse PDF")
            return DocumentAnalysis(
                filename=filename,
                summary="[Error: pdfplumber library not installed]",
                method="local",
            )

        try:
            text_parts = []

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                    # Limit to first 20 pages for analysis
                    if i >= 19:
                        text_parts.append(f"[... {len(pdf.pages) - 20} more pages ...]")
                        break

            full_text = "\n\n".join(text_parts)

            # Create summary from first part of text
            summary = full_text[:500]
            if len(full_text) > 500:
                summary = summary[:497] + "..."

            return DocumentAnalysis(
                filename=filename,
                summary=summary,
                changes=[],  # PDFs don't have tracked changes
                review_items=[],
                questions=[],
                raw_text=full_text,
                method="local",
            )

        except Exception as e:
            logger.error("Failed to parse PDF", filename=filename, error=str(e))
            return DocumentAnalysis(
                filename=filename,
                summary=f"[Error parsing PDF: {str(e)}]",
                method="local",
            )

    async def _parse_xlsx(self, filename: str, content: bytes) -> DocumentAnalysis:
        """Parse XLSX file using openpyxl.

        Extracts sheet info, dimensions, and sample data.
        """
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.error("openpyxl not installed, cannot parse XLSX")
            return DocumentAnalysis(
                filename=filename,
                summary="[Error: openpyxl library not installed]",
                method="local",
            )

        try:
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)

            sheets_info = []
            full_text_parts = []

            for sheet_name in wb.sheetnames[:10]:  # Limit to first 10 sheets
                ws = wb[sheet_name]

                # Get dimensions
                max_row = ws.max_row or 0
                max_col = ws.max_column or 0

                sheets_info.append(f"'{sheet_name}': {max_row} rows x {max_col} cols")

                # Extract headers (first row)
                headers = []
                if max_row > 0:
                    for cell in ws[1]:
                        if cell.value:
                            headers.append(str(cell.value))

                if headers:
                    full_text_parts.append(f"Sheet '{sheet_name}' headers: {', '.join(headers[:10])}")

            wb.close()

            summary = f"Excel workbook with {len(wb.sheetnames)} sheet(s): {'; '.join(sheets_info[:5])}"
            if len(sheets_info) > 5:
                summary += f" (and {len(sheets_info) - 5} more)"

            return DocumentAnalysis(
                filename=filename,
                summary=summary,
                changes=[],
                review_items=[],
                questions=[],
                raw_text="\n".join(full_text_parts),
                method="local",
            )

        except Exception as e:
            logger.error("Failed to parse XLSX", filename=filename, error=str(e))
            return DocumentAnalysis(
                filename=filename,
                summary=f"[Error parsing XLSX: {str(e)}]",
                method="local",
            )


# Singleton instance
_parser: LocalDocumentParser | None = None


def get_document_parser() -> LocalDocumentParser:
    """Get or create the document parser singleton."""
    global _parser
    if _parser is None:
        _parser = LocalDocumentParser()
    return _parser
